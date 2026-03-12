from mcp.types import TextContent, ImageContent
from docx import Document
import pandas as pd
import base64
import os


def _table_to_markdown(table) -> str:
    data = [[cell.text for cell in row.cells] for row in table.rows]
    if not data:
        return ""
    df = pd.DataFrame(data[1:], columns=data[0])
    return df.to_markdown(index=False)


def read_docx(path: str) -> list:
    """Extract text, tables, and images from a Word document."""
    if not os.path.exists(path):
        raise ValueError(f"File not found: {path}")

    try:
        doc = Document(path)
    except Exception as e:
        raise ValueError(f"Could not open DOCX: {e}") from e

    output = []
    texts = []

    for block in doc.element.body:
        tag = block.tag.split("}")[-1]

        if tag == "p":
            for para in doc.paragraphs:
                if para._element is block:
                    line = para.text.strip()
                    if line:
                        texts.append(line)
                    break

        elif tag == "tbl":
            for table in doc.tables:
                if table._element is block:
                    md = _table_to_markdown(table)
                    if md:
                        texts.append(md)
                    break

    if texts:
        output.append(TextContent(type="text", text="\n".join(texts)))

    # Extract inline images
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_part = rel.target_part
            b64 = base64.standard_b64encode(img_part.blob).decode("utf-8")
            mime = img_part.content_type
            output.append(ImageContent(type="image", data=b64, mimeType=mime))

    return output
